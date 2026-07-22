
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level)
    if level == 0:
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return heading

def add_paragraph(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return para

def md_to_docx(md_path, docx_path):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, md_path)
    docx_path = os.path.join(script_dir, docx_path)
    print(f"Processing {md_path} -> {docx_path}")
    if not os.path.exists(md_path):
        print(f"ERROR: File {md_path} does not exist!")
        print(f"Directory contents: {os.listdir(script_dir)}")
        return False
    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    print(f"Read {len(lines)} lines from {md_path}")

    current_text = ""
    for line in lines:
        line = line.rstrip('\n')
        if line.startswith('# '):
            if current_text:
                add_paragraph(doc, current_text.strip())
                current_text = ""
            add_heading(doc, line[2:], 0)
        elif line.startswith('## '):
            if current_text:
                add_paragraph(doc, current_text.strip())
                current_text = ""
            add_heading(doc, line[3:], 1)
        elif line.startswith('### '):
            if current_text:
                add_paragraph(doc, current_text.strip())
                current_text = ""
            add_heading(doc, line[4:], 2)
        elif line.strip() == "":
            if current_text:
                add_paragraph(doc, current_text.strip())
                current_text = ""
        else:
            current_text += line + " "

    # Add any remaining text
    if current_text:
        add_paragraph(doc, current_text.strip())

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")
    return True

if __name__ == "__main__":
    print("Starting conversion...")
    md_to_docx("CS-4107-Corrected.md", "CS-4107-Final.docx")
    md_to_docx("CS-4108-Corrected.md", "CS-4108-Final.docx")
    print("All conversions complete!")



